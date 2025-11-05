"""
文档处理器 - 解析PDF、Word、Excel等格式的财报
"""
from typing import Dict, Any, Optional
from pathlib import Path
import uuid
from datetime import datetime
from loguru import logger
import aiofiles
from backend.core.config import settings
from backend.models.multimodal.chart_parser import ChartParser


class DocumentProcessor:
    """文档处理器 - 提取文本、表格、图表数据"""
    
    def __init__(self):
        self.upload_dir = Path(settings.UPLOAD_DIR)
        self.upload_dir.mkdir(parents=True, exist_ok=True)
    
    async def save_uploaded_file(self, file) -> Dict[str, Any]:
        """保存上传的文件"""
        file_id = str(uuid.uuid4())
        file_ext = "." + file.filename.split(".")[-1].lower()
        file_path = self.upload_dir / f"{file_id}{file_ext}"
        
        # 异步保存文件
        async with aiofiles.open(file_path, 'wb') as f:
            content = await file.read()
            await f.write(content)
        
        return {
            "file_id": file_id,
            "path": str(file_path),
            "filename": file.filename,
            "size": len(content),
        }
    
    async def process_document(
        self,
        file_path: str,
        file_type: str,
        company: Optional[str] = None,
        year: Optional[int] = None,
        report_type: str = "annual",
    ) -> Dict[str, Any]:
        """
        处理文档，提取结构化数据
        
        Returns:
            提取的数据，包括文本、表格、指标等
        """
        try:
            if file_type == ".pdf":
                return await self._process_pdf(file_path, company, year, report_type)
            elif file_type in [".docx", ".doc"]:
                return await self._process_word(file_path, company, year, report_type)
            elif file_type in [".xlsx", ".xls"]:
                return await self._process_excel(file_path, company, year, report_type)
            else:
                raise ValueError(f"不支持的文件格式: {file_type}")
                
        except Exception as e:
            logger.error(f"文档处理失败: {e}")
            raise
    
    async def _process_pdf(self, file_path: str, company: Optional[str], year: Optional[int], report_type: str) -> Dict[str, Any]:
        """处理PDF文件"""
        try:
            import pdfplumber
            
            extracted_data = {
                "text": "",
                "tables": [],
                "metadata": {
                    "company": company,
                    "year": year,
                    "report_type": report_type,
                },
            }
            
            with pdfplumber.open(file_path) as pdf:
                # 提取文本
                texts = []
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        texts.append(text)
                
                extracted_data["text"] = "\n".join(texts)
                
                # 提取表格
                for page in pdf.pages:
                    tables = page.extract_tables()
                    for table in tables:
                        if table:
                            extracted_data["tables"].append(table)
            
            # 提取关键指标（使用OCR和NLP，这里简化）
            extracted_data["indicators"] = await self._extract_indicators(extracted_data["text"])
            
            logger.info(f"PDF处理完成: {len(extracted_data['text'])} 字符, {len(extracted_data['tables'])} 表格")
            
            return extracted_data
            
        except Exception as e:
            logger.error(f"PDF处理失败: {e}")
            raise
    
    async def _process_word(self, file_path: str, company: Optional[str], year: Optional[int], report_type: str) -> Dict[str, Any]:
        """处理Word文件"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            
            extracted_data = {
                "text": "\n".join([para.text for para in doc.paragraphs]),
                "tables": [],
                "metadata": {
                    "company": company,
                    "year": year,
                    "report_type": report_type,
                },
            }
            
            # 提取表格
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                extracted_data["tables"].append(table_data)
            
            extracted_data["indicators"] = await self._extract_indicators(extracted_data["text"])
            
            return extracted_data
            
        except Exception as e:
            logger.error(f"Word处理失败: {e}")
            raise
    
    async def _process_excel(self, file_path: str, company: Optional[str], year: Optional[int], report_type: str) -> Dict[str, Any]:
        """处理Excel文件"""
        try:
            import pandas as pd
            
            # 读取所有工作表
            excel_file = pd.ExcelFile(file_path)
            sheets_data = {}
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                sheets_data[sheet_name] = df.to_dict('records')
            
            extracted_data = {
                "text": "",  # Excel主要是表格数据
                "tables": list(sheets_data.values()),
                "sheet_names": list(sheets_data.keys()),
                "metadata": {
                    "company": company,
                    "year": year,
                    "report_type": report_type,
                },
            }
            
            # 从表格中提取指标
            extracted_data["indicators"] = await self._extract_indicators_from_tables(extracted_data["tables"])
            
            return extracted_data
            
        except Exception as e:
            logger.error(f"Excel处理失败: {e}")
            raise
    
    async def _extract_indicators(self, text: str) -> Dict[str, Any]:
        """从文本中提取财务指标（简化版，实际应该使用NER模型）"""
        # 这里应该使用NER模型或规则引擎
        indicators = {}
        
        # 简化的关键词匹配
        indicator_keywords = {
            "营收": ["营业收入", "营收", "收入"],
            "净利润": ["净利润", "归母净利润"],
            "ROE": ["净资产收益率", "ROE"],
            "不良率": ["不良贷款率", "不良率"],
        }
        
        for indicator, keywords in indicator_keywords.items():
            for keyword in keywords:
                if keyword in text:
                    # 提取数值（简化版）
                    import re
                    pattern = rf"{keyword}.*?([\d,]+\.?\d*)\s*[万亿万千百十]?元"
                    matches = re.findall(pattern, text)
                    if matches:
                        indicators[indicator] = matches[0]
        
        return indicators
    
    async def _extract_indicators_from_tables(self, tables: list) -> Dict[str, Any]:
        """从表格中提取指标"""
        indicators = {}
        
        # 遍历所有表格，查找指标
        for table in tables:
            if not table:
                continue
            
            # 查找表头中的指标名称
            for row in table[:10]:  # 检查前10行
                if isinstance(row, list):
                    for i, cell in enumerate(row):
                        if isinstance(cell, str):
                            # 检查是否是指标名称
                            if any(keyword in cell for keyword in ["营收", "净利润", "ROE", "不良率"]):
                                # 尝试提取对应的数值
                                if i + 1 < len(row):
                                    indicators[cell] = row[i + 1]
        
        return indicators
    
    async def get_processing_status(self, file_id: str) -> Dict[str, Any]:
        """获取文件处理状态"""
        # 这里应该从数据库查询处理状态
        return {
            "file_id": file_id,
            "status": "completed",  # pending, processing, completed, failed
            "progress": 100,
        }

